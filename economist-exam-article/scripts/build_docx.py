#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_docx.py — 将 economist-exam-article 技能的内容 JSON 渲染为排版规范的 .docx 文件

用法:
    python3 build_docx.py input.json -o output.docx

JSON 格式（详见 SKILL.md）:
{
  "title": "英文标题",
  "subtitle": "中文副标题",
  "section": "Business",
  "english_article": "段落1\\n\\n段落2……（**word**(*中文释义*) 内联标记）",
  "chinese_article": "中文译文（段落之间空行分隔）",
  "questions": [
    {"number": 21, "stem": "题干", "options": ["A. ...", "B. ...", "C. ...", "D. ..."],
     "answer": "B", "explanation": "解析"}
  ],
  "vocabulary": [{"word": "ephemeral", "pos": "adj.", "translation": "短暂的"}]
}

输出排版（四部分，分页符分隔）:
  Part 1 练习页: 标题区 → 纯英文原文(无加粗无翻译) → 阅读理解题(21-25)
  Part 2 解析页: 参考答案与解析(速查 + 逐题解析)
  Part 3 精读页: 带注释英文原文(目标词加粗+斜体括号释义) → 中文翻译
  Part 4 默写页: 词汇表(Word/词性/中文释义 —— 释义列为空供默写)
"""

import argparse
import json
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ----------------------------------------------------------------------------
# 常量
# ----------------------------------------------------------------------------
LATIN_FONT = "Times New Roman"   # 西文字体
EA_FONT = "宋体"                 # 正文中文字体
EA_HEADING_FONT = "黑体"         # 标题中文字体

BODY_SIZE = 11                   # 正文字号(pt)
TITLE_SIZE = 16
SUBTITLE_SIZE = 12
HEADING_SIZE = 14
META_SIZE = 10

# 字符间距（w:spacing val，单位 1/20 磅）：轻微放宽让排版更透气
CHAR_SPACING = 4                 # 正文 0.2pt
CHAR_SPACING_HEADING = 6         # 标题 0.3pt

ACCENT_COLOR = RGBColor(0x1F, 0x3A, 0x5F)  # 深蓝，用于分区标题
ANSWER_COLOR = RGBColor(0x8B, 0x1A, 0x1A)  # 深红，用于答案
GRAY_COLOR = RGBColor(0x59, 0x59, 0x59)

# 匹配 **加粗** 或 (*斜体括号*) 内联标记（按出现顺序切分）
MARKER_RE = re.compile(r"(\*\*.+?\*\*|\(\*.*?\*\))")


# ----------------------------------------------------------------------------
# 基础工具
# ----------------------------------------------------------------------------
def set_font(run, size=BODY_SIZE, bold=False, italic=False, color=None,
             ea=EA_FONT, latin=LATIN_FONT, spacing=CHAR_SPACING):
    """设置 run 的字体属性（Times New Roman + 宋体；可调字符间距 w:spacing）。"""
    run.font.name = latin
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), ea)
    if spacing is not None:
        # 字符间距：w:spacing val，单位 1/20 磅（与段落行距 w:spacing 是不同元素）
        spacing_el = rPr.find(qn("w:spacing"))
        if spacing_el is None:
            spacing_el = rPr.makeelement(qn("w:spacing"), {})
            rPr.append(spacing_el)
        spacing_el.set(qn("w:val"), str(spacing))


def setup_document_defaults(doc):
    """设置 Normal 样式默认字体为 Times New Roman + 宋体，确保全文档统一。"""
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(BODY_SIZE)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), EA_FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)


def strip_markers(text):
    """从带标记文本剥离出纯文本：**word**(*译*) -> word（丢弃翻译与星号）。"""
    out = []
    for token in MARKER_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            out.append(token[2:-2])
        elif token.startswith("(*") and token.endswith("*)"):
            continue  # 翻译括号在纯阅读版中省略
        else:
            out.append(token)
    return "".join(out)


def add_marked_runs(paragraph, text, size=BODY_SIZE):
    """解析 **word**(*翻译*) 标记写入 runs：
       **word** -> 加粗；(*翻译*) -> 斜体并保留括号。其余文本原样。"""
    for token in MARKER_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=size, bold=True, italic=False)
        elif token.startswith("(*") and token.endswith("*)"):
            run = paragraph.add_run(f"({token[2:-2]})")  # 去掉星号，保留括号，整体斜体
            set_font(run, size=size, bold=False, italic=True)
        else:
            run = paragraph.add_run(token)
            set_font(run, size=size, bold=False, italic=False)
    return paragraph


def add_heading(doc, text):
    """分区标题（宋体加粗深蓝，如【Part 1 · 英文原文】）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=HEADING_SIZE, bold=True, color=ACCENT_COLOR,
             ea=EA_FONT, spacing=CHAR_SPACING_HEADING)
    return p


def add_body_paragraph(doc, text, size=BODY_SIZE, indent=True, plain=False):
    """正文段落：首行缩进 2 字符、两端对齐、1.5 倍行距。
       plain=True 时按纯文本输出（剥离标记）。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if indent:
        pf.first_line_indent = Pt(size * 2)
    pf.line_spacing = 1.5
    pf.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if plain:
        set_font(p.add_run(strip_markers(text)), size=size)
    else:
        add_marked_runs(p, text, size=size)
    return p


# ----------------------------------------------------------------------------
# 文档构建
# ----------------------------------------------------------------------------
def build(data, out_path):
    doc = Document()
    setup_document_defaults(doc)  # 全局默认字体：Times New Roman + 宋体

    # 页面设置：A4，标准页边距
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.5)

    title = data.get("title", "Untitled")
    subtitle = data.get("subtitle", "")
    section = data.get("section", "")
    eng_article = data.get("english_article", "")
    chi_article = data.get("chinese_article", "")
    questions = data.get("questions", [])
    vocabulary = data.get("vocabulary", [])

    eng_paras = [p.strip() for p in eng_article.split("\n\n") if p.strip()]
    chi_paras = [p.strip() for p in chi_article.split("\n\n") if p.strip()]

    # ---- 标题区 ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(title), size=TITLE_SIZE, bold=True, ea=EA_HEADING_FONT,
             spacing=CHAR_SPACING_HEADING)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        set_font(p.add_run(subtitle), size=SUBTITLE_SIZE, bold=False, ea=EA_FONT)

    word_count = sum(len(strip_markers(par).split()) for par in eng_paras)
    meta = f"板块：{section or '未指定'} ｜ 英文词数：约 {word_count} 词 ｜ 目标词汇：{len(vocabulary)} 个"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    set_font(p.add_run(meta), size=META_SIZE, italic=True, color=GRAY_COLOR)

    # ================= Part 1 · 练习区 =================
    add_heading(doc, "【Part 1 · 英文原文】（纯阅读版）")
    for par in eng_paras:
        add_body_paragraph(doc, par, plain=True)

    if questions:
        # 阅读理解题与英文原文分开页面
        doc.add_page_break()
        add_heading(doc, "【Part 1 · 阅读理解题】（Text 1 · Q21–Q25）")
        for q in questions:
            num = q.get("number", "")
            stem = q.get("stem", "")
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            set_font(p.add_run(f"{num}. {stem}"), size=BODY_SIZE, bold=True)

            for opt in q.get("options", []):
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Cm(1.0)
                op.paragraph_format.right_indent = Cm(1.0)
                op.paragraph_format.space_after = Pt(2)
                set_font(op.add_run(opt), size=BODY_SIZE)

    # ================= Part 2 · 解析页 =================
    if questions:
        doc.add_page_break()
        add_heading(doc, "【Part 2 · 参考答案与解析】")
        key_line = "　".join(f"{q.get('number')}.{q.get('answer')}" for q in questions)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        set_font(p.add_run("答案速查："), size=BODY_SIZE, bold=True)
        set_font(p.add_run(key_line), size=BODY_SIZE, bold=True, color=ANSWER_COLOR)

        for q in questions:
            num = q.get("number", "")
            ans = q.get("answer", "")
            exp = q.get("explanation", "")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_font(p.add_run(f"{num}."), size=BODY_SIZE, bold=True, color=ANSWER_COLOR)
            set_font(p.add_run(f"【答案】{ans}  "), size=BODY_SIZE, bold=True, color=ANSWER_COLOR)
            set_font(p.add_run(f"【解析】{exp}"), size=BODY_SIZE, ea=EA_FONT)

    # ================= Part 3 · 精读页 =================
    doc.add_page_break()
    add_heading(doc, "【Part 3 · 英文原文】（精读版：目标词加粗 + 释义）")
    for par in eng_paras:
        add_body_paragraph(doc, par)

    if chi_paras:
        add_heading(doc, "【Part 3 · 中文翻译】")
        for par in chi_paras:
            add_body_paragraph(doc, par, size=10.5)

    # ================= Part 4 · 默写页 =================
    if vocabulary:
        doc.add_page_break()
        add_heading(doc, "【Part 4 · 词汇表】（默写版：请填写中文释义）")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, text in enumerate(("Word", "词性", "中文释义")):
            hdr[i].text = ""
            set_font(hdr[i].paragraphs[0].add_run(text), size=10.5, bold=True, ea=EA_FONT)

        for item in vocabulary:
            row = table.add_row().cells
            word = item.get("word", "")
            pos = item.get("pos", "")
            set_font(row[0].paragraphs[0].add_run(word), size=10.5, bold=True)
            set_font(row[1].paragraphs[0].add_run(pos), size=10.5)
            # 第三列（中文释义）留空，供默写

        # 表格列宽
        widths = (Cm(4.5), Cm(2.0), Cm(9.5))
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = width

    doc.save(out_path)
    print(f"OK: {out_path}")


# ----------------------------------------------------------------------------
# 入口
# ----------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Generate Economist-style exam practice docx")
    parser.add_argument("input", help="输入 JSON 文件路径")
    parser.add_argument("-o", "--output", default="output.docx", help="输出 docx 路径")
    args = parser.parse_args()

    try:
        with open(args.input, "r", encoding="utf-8") as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"ERROR: 找不到输入文件 {args.input}", file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"ERROR: JSON 解析失败 - {e}", file=sys.stderr)
        sys.exit(1)

    build(data, args.output)


if __name__ == "__main__":
    main()
